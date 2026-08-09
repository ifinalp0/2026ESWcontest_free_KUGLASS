#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def iter_blocks(parent):
    if isinstance(parent, _Document):
        body = parent.element.body
    elif isinstance(parent, _Cell):
        body = parent._tc
    else:
        raise TypeError(type(parent))
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def emu(value):
    return None if value is None else int(value)


def paragraph_record(p: Paragraph):
    return {
        "kind": "paragraph",
        "style": p.style.name if p.style else None,
        "text": p.text,
        "alignment": str(p.alignment),
        "runs": [
            {
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "size_pt": None if r.font.size is None else r.font.size.pt,
                "font": r.font.name,
            }
            for r in p.runs
        ],
    }


def main(path: str, out: str):
    doc = Document(path)
    blocks = []
    table_summaries = []
    styles = Counter()
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            rec = paragraph_record(block)
            blocks.append(rec)
            styles[rec["style"]] += 1
        else:
            rows = []
            for row in block.rows:
                row_cells = []
                for cell in row.cells:
                    row_cells.append("\n".join(p.text for p in cell.paragraphs))
                    for p in cell.paragraphs:
                        styles[p.style.name if p.style else None] += 1
                rows.append(row_cells)
            table_summaries.append({
                "rows": len(block.rows),
                "cols": len(block.columns),
                "style": block.style.name if block.style else None,
                "grid_cols_emu": [int(gc.w) for gc in block._tbl.tblGrid.gridCol_lst],
                "data": rows,
            })
            blocks.append({"kind": "table", "index": len(table_summaries) - 1})

    sections = []
    for i, s in enumerate(doc.sections):
        sections.append({
            "index": i,
            "page_width_emu": emu(s.page_width),
            "page_height_emu": emu(s.page_height),
            "top_margin_emu": emu(s.top_margin),
            "bottom_margin_emu": emu(s.bottom_margin),
            "left_margin_emu": emu(s.left_margin),
            "right_margin_emu": emu(s.right_margin),
            "header_distance_emu": emu(s.header_distance),
            "footer_distance_emu": emu(s.footer_distance),
            "start_type": str(s.start_type),
            "header": [p.text for p in s.header.paragraphs],
            "footer": [p.text for p in s.footer.paragraphs],
        })

    style_defs = {}
    for name in styles:
        if not name:
            continue
        style = doc.styles[name]
        pf = style.paragraph_format
        style_defs[name] = {
            "type": str(style.type),
            "font": style.font.name,
            "font_size_pt": None if style.font.size is None else style.font.size.pt,
            "bold": style.font.bold,
            "italic": style.font.italic,
            "alignment": str(pf.alignment),
            "space_before_pt": None if pf.space_before is None else pf.space_before.pt,
            "space_after_pt": None if pf.space_after is None else pf.space_after.pt,
            "line_spacing": str(pf.line_spacing),
            "left_indent_emu": emu(pf.left_indent),
            "first_line_indent_emu": emu(pf.first_line_indent),
        }

    rel_types = Counter(rel.reltype for rel in doc.part.rels.values())
    payload = {
        "path": str(Path(path).resolve()),
        "sections": sections,
        "style_usage": styles,
        "style_defs": style_defs,
        "blocks": blocks,
        "tables": table_summaries,
        "inline_shapes": len(doc.inline_shapes),
        "relationships": rel_types,
        "core_properties": {
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "author": doc.core_properties.author,
            "last_modified_by": doc.core_properties.last_modified_by,
        },
    }
    Path(out).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
