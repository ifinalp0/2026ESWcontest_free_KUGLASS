#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor, Twips


REFERENCE = Path(
    "/Users/mooyoung/Documents/건국대학교 2-1/DECA/쿠글(KUGLASS)/"
    "쿠글_외/임소경 개발계획서/임소경 개발 계획서.docx"
)
EXPECTED_SHA256 = "b867eb9fe841c78151aa72047aa3b685e05637da181c847695cbbc6e187dc0f9"
OUTPUT = Path(sys.argv[1]).resolve()

FONT_BODY = "나눔명조"
FONT_SANS = "나눔고딕"
COLOR_TEXT = "111111"
COLOR_GREEN = "087A2B"
COLOR_GREEN_LIGHT = "D9F2DE"
COLOR_BLUE_HEADER = "D9E5F2"
COLOR_BLUE_BORDER = "AFC4D8"
COLOR_BLUE_DARK = "58728A"
COLOR_GRAY = "666666"
COLOR_GRAY_LIGHT = "F3F5F7"
COLOR_WHITE = "FFFFFF"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_run_font(run, name=FONT_BODY, size=10.5, bold=None, color=COLOR_TEXT, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name, size, bold=None, color=COLOR_TEXT):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips):
    cell.width = Twips(width_twips)
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_twips))
    tcw.set(qn("w:type"), "dxa")


def set_table_borders(table, outer=8, inner=5, color=COLOR_BLUE_BORDER):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(outer if edge in {"top", "left", "bottom", "right"} else inner))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def configure_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
    set_table_borders(table)


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trpr.append(cant)


def set_paragraph_border(paragraph, edge, color, size=12, space=6):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    node = pbdr.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        pbdr.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))
    node.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def create_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level, marker, left, hanging in ((0, "•", 360, 180), (1, "–", 720, 180), (2, "·", 1080, 180)):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "bullet")
        lvl.append(numfmt)
        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), marker)
        lvl.append(lvltext)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        ppr.append(ind)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), FONT_BODY)
        rfonts.set(qn("w:hAnsi"), FONT_BODY)
        rfonts.set(qn("w:eastAsia"), FONT_BODY)
        rpr.append(rfonts)
        lvl.append(rpr)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id, level=0):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def setup_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    set_style_font(normal, FONT_BODY, 10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.widow_control = True

    def ensure(name, size, bold, before, after, keep=True, color=COLOR_TEXT):
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, FONT_BODY, size, bold, color)
        style.base_style = normal
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.1
        pf.keep_with_next = keep
        pf.widow_control = True
        return style

    ensure("KUG Section", 15, True, 10, 5)
    ensure("KUG Subsection", 11.5, True, 7, 3)
    ensure("KUG Minor", 10.5, True, 5, 2)
    ensure("KUG Caption", 9, False, 2, 2, color=COLOR_GRAY)
    ensure("KUG Reference", 9.5, False, 1, 2, keep=False)
    ensure("KUG Bullet", 10.5, False, 0, 2, keep=False)
    return styles


def add_section_heading(document, text, page_break=False):
    if page_break:
        document.add_page_break()
    p = document.add_paragraph(style="KUG Section")
    p.add_run(f"□ {text}")
    return p


def add_subheading(document, text):
    p = document.add_paragraph(style="KUG Subsection")
    p.add_run(f"○ {text}")
    return p


def add_minor_heading(document, text):
    p = document.add_paragraph(style="KUG Minor")
    p.add_run(text)
    return p


def add_body(document, text, bold_prefix=None):
    p = document.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(document, text, num_id, level=0, bold_prefix=None):
    p = document.add_paragraph(style="KUG Bullet")
    apply_bullet(p, num_id, level)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(document, label, text, fill=COLOR_GRAY_LIGHT, border=COLOR_BLUE_BORDER):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Twips(140)
    p.paragraph_format.right_indent = Twips(140)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    shade_paragraph(p, fill)
    for edge in ("top", "left", "bottom", "right"):
        set_paragraph_border(p, edge, border, size=6, space=6)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, FONT_SANS, 9.5, True, COLOR_BLUE_DARK)
    r2 = p.add_run(text)
    set_run_font(r2, FONT_BODY, 10.0, False, COLOR_TEXT)
    return p


def add_table(document, headers, rows, widths, font_size=9.5, header_font_size=9.5,
              status_col=None, status_colors=None, first_col_bold=False):
    table = document.add_table(rows=1, cols=len(headers))
    configure_table_geometry(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    prevent_row_split(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, COLOR_BLUE_HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(header)
        set_run_font(r, FONT_BODY, header_font_size, True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if status_col is not None and index == status_col and status_colors:
                fill = status_colors.get(str(value))
                if fill:
                    set_cell_shading(cell, fill)
            parts = value if isinstance(value, list) else [str(value)]
            cell.text = ""
            for part_index, part in enumerate(parts):
                p = cell.paragraphs[0] if part_index == 0 else cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
                if status_col is not None and index == status_col:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.08
                r = p.add_run(part)
                set_run_font(r, FONT_BODY, font_size, bool(first_col_bold and index == 0))
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
    document.add_paragraph(style="KUG Caption")
    return table


def add_summary_table(document, widths):
    rows = [
        ("분야", "자동차/모빌리티"),
        ("팀명", "쿠글"),
        ("작품명", "능동형 스마트 글라스 모빌리티(KUGLASS)"),
        ("작품 범위", "1:10 차량 모형용 PDLC 4채널 시연 프로토타입"),
        ("현재 단계", "제품 코드·통신 계약·제작 PCB 기준 확립 및 비실기 검증 완료 / 통합 HIL·광학·열·전력 실측 진행 전"),
        ("작품 설명", [
            "카메라 1대와 DS18B20 내부온도센서 1개의 입력을 ESP32_A가 분석하여 CH0~CH3의 목표 MI를 산출한다.",
            "ESP32_B는 20 Hz full-frame 명령을 검증하고 Logic Carrier와 Power Stage PCB 4장을 통해 4채널 SPWM 출력을 적용한다.",
            "열부하 차광, 차박·주차 프라이버시, 카메라 강광 완화를 하나의 태블릿 HMI에서 시연한다.",
        ]),
    ]
    table = document.add_table(rows=0, cols=2)
    configure_table_geometry(table, widths)
    for label, value in rows:
        row = table.add_row()
        prevent_row_split(row)
        left, right = row.cells
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, "E7E7E7")
        pl = left.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_after = Pt(0)
        rl = pl.add_run(label)
        set_run_font(rl, FONT_BODY, 9.5, True)
        values = value if isinstance(value, list) else [value]
        right.text = ""
        for i, item in enumerate(values):
            p = right.paragraphs[0] if i == 0 else right.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(item)
            set_run_font(r, FONT_BODY, 9.5)
            if isinstance(value, list):
                apply_bullet(p, bullet_num_id, 0)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell, top=90, bottom=90)
    document.add_paragraph(style="KUG Caption")
    return table


def set_document_compatibility(document):
    settings = document.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    mode = compat.find(qn("w:compatSetting"))
    if mode is None:
        mode = OxmlElement("w:compatSetting")
        compat.append(mode)
    mode.set(qn("w:name"), "compatibilityMode")
    mode.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    mode.set(qn("w:val"), "15")


if sha256(REFERENCE) != EXPECTED_SHA256:
    raise RuntimeError("Reference DOCX hash mismatch; fresh distillation required")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(REFERENCE, OUTPUT)
doc = Document(OUTPUT)

# Retain the source section properties but replace the obsolete report body.
body = doc._element.body
sect_pr = body.sectPr
for child in list(body):
    if child is not sect_pr:
        body.remove(child)

section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Emu(7_560_310)
section.page_height = Emu(10_692_130)
section.left_margin = Emu(720_090)
section.right_margin = Emu(720_090)
section.top_margin = Emu(698_500)
section.bottom_margin = Emu(539_750)
section.header_distance = Emu(158_750)
section.footer_distance = Emu(0)

setup_styles(doc)
set_document_compatibility(doc)
bullet_num_id = create_bullet_numbering(doc)

doc.core_properties.title = "교내 창의설계경진대회 예선 보고서"
doc.core_properties.subject = "KUGLASS 현재 개발 현황 및 잔여 검증 계획"
doc.core_properties.author = "쿠글"
doc.core_properties.last_modified_by = "쿠글"
doc.core_properties.keywords = "KUGLASS, PDLC, ESP32-S3, 스마트 글라스, 창의설계"

usable_twips = int(
    (section.page_width - section.left_margin - section.right_margin) / 635
)
table_width = usable_twips - 120

# Opening title block.
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(8)
title.paragraph_format.space_after = Pt(3)
set_paragraph_border(title, "top", COLOR_GREEN, size=20, space=10)
r = title.add_run("교내 창의설계경진대회 예선 보고서")
set_run_font(r, FONT_BODY, 23, True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(4)
set_paragraph_border(subtitle, "bottom", COLOR_GREEN_LIGHT, size=18, space=8)
r = subtitle.add_run("[자동차/모빌리티]")
set_run_font(r, FONT_BODY, 13, True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
meta.paragraph_format.space_before = Pt(0)
meta.paragraph_format.space_after = Pt(6)
r = meta.add_run("KUGLASS  |  작성 기준 2026.08.10  |  기준 커밋 839666a4ca88")
set_run_font(r, FONT_SANS, 8.5, False, COLOR_GRAY)

add_section_heading(doc, "개발 개요")
add_subheading(doc, "요약 설명")
add_summary_table(doc, [1450, table_width - 1450])

add_subheading(doc, "개발 배경 및 동기")
for text in [
    "차량 유리는 외부 환경을 차단하면서 운전자와 승객의 시야를 제공하는 핵심 인터페이스이다. 고정 투과율 유리나 수동 가림막은 주행, 정차, 주차, 차박처럼 요구가 빠르게 바뀌는 상황에 능동적으로 대응하기 어렵다.",
    "강한 역광과 직사광은 운전자에게 불편을 주고 카메라의 포화 영역을 확대할 수 있다. 하절기 일사로 인한 차량 내부 열부하는 냉방 부담과 승차 불쾌감을 높이며, 주차·차박 상황에서는 외부 시선 차단이 필요하다.",
    "KUGLASS는 PDLC의 산란 정도를 연속적으로 바꿀 수 있다는 특성을 이용해, 카메라와 내부온도에서 얻은 상황 정보를 네 개 채널의 목표 MI로 변환하고 실제 전력 구동부까지 폐루프 상태를 확인하는 모형 시연 시스템을 제안한다.",
]:
    add_bullet(doc, text, bullet_num_id)

add_subheading(doc, "개발 목적과 제품 범위")
for prefix, text in [
    ("목표 1. ", "ESP32_A에서 카메라·내부온도 입력 품질을 판단하고 CH0~CH3의 권위 있는 target_mi를 산출한다."),
    ("목표 2. ", "ESP32_B에서 strict full-frame 검증, 4채널 16 kHz/60 Hz SPWM, 방향 blanking, Fault·E-Stop·TTL safe-off를 수행한다."),
    ("목표 3. ", "MacBook TabUI와 태블릿 HMI에서 센서, 제어 출처, target/commanded/applied MI, A/B 연결 상태와 Fault를 구분하여 시연한다."),
    ("범위. ", "본 작품은 1:10 차량 모형용 시연 프로토타입이며 실차 안전 장치 또는 자율주행 인지 성능 향상 장치로 주장하지 않는다."),
]:
    add_bullet(doc, prefix + text, bullet_num_id, bold_prefix=prefix)

add_subheading(doc, "현재 시스템 구성")
add_callout(
    doc,
    "데이터·제어 흐름",
    "카메라 + DS18B20 내부온도 → ESP32_A(입력·정책·목표 MI) → 외부 3선 UART → ESP32_B(검증·SPWM·로컬 차단) → Logic Carrier → Power Stage PCB ×4 → PDLC CH0~CH3",
    fill="EEF7F0",
    border="A8C9B0",
)

component_rows = [
    ("카메라·온도", "OV2640 1대, YwRobot SEN050007 DS18B20 1개", "실제 제품 입력. 결측·stale을 상태로 관리"),
    ("ESP32_A", "센서 처리, 정책, LUT/MI servo, 목표 MI", "AUTO 권위 보유, B로 20 Hz full frame 전송"),
    ("ESP32_B", "4채널 SPWM, 명령 검증, Fault·TTL 차단", "목표를 재계산하지 않고 실제 applied_mi 회신"),
    ("Logic Carrier", "E-Stop enable gate, Fault pull-up, ADC filter, J7", "ESP32_B를 U3에 장착하며 우회 직결 금지"),
    ("Power Stage ×4", "단일 채널 H-Bridge + LC filter + RUN_OK", "동일 보드를 CH0~CH3에 한 장씩 배치"),
    ("TabUI", "MacBook 백엔드, 브라우저 HMI, USB gateway", "LIVE/MOCK/REPLAY 분리, LIVE 자동 MI 미계산"),
]
add_table(doc, ["구성요소", "주요 구성·기능", "책임 경계"], component_rows,
          [1600, 3350, table_width - 4950], font_size=9.2, first_col_bold=True)

add_subheading(doc, "채널 정의와 MI 의미")
channel_rows = [
    ("CH0", "운전석 창문", "운전석측 ROI", "독립 목표·적용 상태"),
    ("CH1", "조수석 창문 + 선루프", "조수석측 ROI", "두 PDLC 영역이 하나의 MI 공유"),
    ("CH2", "운전석 옆 창문", "운전석측 ROI", "thermal/privacy와 카메라 보조"),
    ("CH3", "조수석 옆 창문", "조수석측 ROI", "thermal/privacy와 카메라 보조"),
]
add_table(doc, ["채널", "PDLC 영역", "카메라 입력", "비고"], channel_rows,
          [850, 2650, 2200, table_width - 5700], font_size=9.2, first_col_bold=True)
add_callout(doc, "MI 기준", "MI가 클수록 CLEAR 방향이며 현재 운용 상한 0.60을 완전 투명으로 취급한다. MI 0.0 또는 disable은 enable 해제와 강산란 방향의 safe-off이며, 전원 rail 전체 차단을 뜻하지 않는다.")

# Development direction and strategy.
add_section_heading(doc, "개발 방향 및 전략", page_break=True)
add_subheading(doc, "개발 방향")
direction_items = [
    ("책임 분리", "센서·정책·목표값은 ESP32_A, 출력·로컬 안전은 ESP32_B, 화면·명령 중계는 TabUI가 맡는다. UI 장애가 발생해도 ESP32_A의 AUTO와 A→B heartbeat가 유지되는 구조다."),
    ("4채널 현행 구성", "채널은 항상 CH0~CH3이며, 제작된 Logic Carrier 1장과 동일한 단일 채널 Power Stage PCB 4장을 사용한다. 이전 8채널 계획은 현재 제품 범위가 아니다."),
    ("계측 기반 완성", "소프트웨어 상수와 회로 명목값을 실측값으로 포장하지 않는다. MI–Vrms–투과도 LUT, 카메라 개선량, 내부온도 기준, 보드별 전류·온도 변환은 HIL 계측으로 확정한다."),
    ("안전 우선 통합", "부팅·invalid frame·통신 timeout·watchdog·E-Stop은 전체 safe-off, latched Power Stage Fault는 해당 채널만 safe-off한다. 72 V와 PDLC는 logic-only·저전압 검증 이후에 연결한다."),
    ("시연 정직성", "LIVE, MOCK, REPLAY를 명시적으로 구분하고 LIVE 단절 시 마지막 실제 상태를 STALE/OFFLINE으로 표시한다. MOCK으로 자동 전환하거나 baseline 없는 개선율을 생성하지 않는다."),
]
for prefix, text in direction_items:
    add_bullet(doc, f"{prefix}: {text}", bullet_num_id, bold_prefix=f"{prefix}: ")

add_subheading(doc, "소프트웨어 개발 방법 및 활용 기술")
software_sections = [
    ("카메라 입력·영상 지표", "ESP32 카메라 드라이버로 VGA 640×480 RGB565 frame을 수집한다. 장착 방향의 180° 보정을 capture 단계에서 적용하고, 화면 왼쪽/오른쪽 ROI에서 밝기·포화·highlight·Edge Density를 계산한다. AE exposure/gain 단위가 검증되기 전에는 ae_metadata_valid=false로 유지한다."),
    ("내부온도 입력", "GPIO41의 DS18B20을 OneWire로 읽고 CRC, 값 범위, 5초 stale 여부를 검사한다. 결측값을 임의 기본 온도처럼 사용하지 않으며 MOCK/HIL의 합성 온도는 물리 센서와 구분한다."),
    ("정책·MI servo", "주행·정차·차박·주차 모드, privacy, thermal risk와 카메라 강광 지표를 결합해 target_mi를 만든다. AUTO에는 0.01 MI deadband와 rate limit을 적용하고 강광 fast-attack 경로를 둔다. 관리자 지속/일반 TTL 수동 제어는 AUTO와 구분된다."),
    ("A–B 통신", "ESP32_A는 v=1 actuator_command를 20 Hz로 보내며 CH0~CH3를 정확히 한 번씩 포함한다. 기본 TTL은 250 ms다. ESP32_B는 stale/duplicate sequence, 누락·중복 채널, 범위 밖 MI, malformed JSON을 전체 거부한다."),
    ("ESP32_B 출력·상태", "MCPWM 기반 16 kHz carrier와 60 Hz polarity로 Simplified Unipolar SPWM을 생성한다. 정상 변화에는 감소 12.0 MI/s, 증가 4.0 MI/s의 최종 slew를 적용하고 100 ms 주기로 applied_mi·Fault·ADC 상태를 회신한다."),
    ("TabUI·HMI", "React·TypeScript·Vite·React Three Fiber 기반 HMI와 Python 백엔드가 MacBook에서 실행된다. DevKit USB Serial/JTAG CDC를 통해 ESP32_A와 연결하며 LIVE/MOCK/REPLAY, 관리자 콘솔, 카메라 Evidence View와 on-demand JPEG를 제공한다."),
]
for heading, text in software_sections:
    add_minor_heading(doc, f"- {heading}")
    add_body(doc, text)

add_subheading(doc, "하드웨어 개발 방법 및 안전 구조")
hardware_items = [
    ("센서·정책 장치", "ESP32_A DevKit에 카메라 1대와 3.3 V 외부전원의 DS18B20 모듈 1개를 연결한다. A–B UART(GPIO39/40)와 DS18B20 GPIO41, 카메라 GPIO4~18의 물리 header 충돌을 실기 전에 확인한다."),
    ("A–B 외부 harness", "A GPIO39 TX → B GPIO44 RX, A GPIO40 RX ← B GPIO43 TX, 공통 GND의 3선 교차 연결을 사용한다. Logic Carrier에는 이 UART가 라우팅되지 않으며 ESP32_B DevKit bridge contention을 실측한다."),
    ("Logic Carrier", "ESP32_B를 U3에 장착한다. U4가 CHx_ENABLE = EN_GLOBAL AND ENABLE_CHx를 만들고 J7이 PWM_MAG, DIR, enable, FAULT_N, ADC, 3.3 V, 12 V를 각 Power Stage에 분배한다. 모든 짝수 핀은 GND다."),
    ("Power Stage", "동일한 단일 채널 보드 4장을 CH0~CH3에 사용한다. 각 보드는 H-Bridge, IRS2104 driver, L1/L2 470 µH와 C10 470 nF filter, RUN_OK = CHx_ENABLE AND FAULT_N 구조를 가진다."),
    ("보호·계측 경계", "FAULT_N falling edge는 해당 채널을 즉시 차단하고, E-Stop falling edge는 네 채널을 즉시 차단한다. ADC 8채널 raw/mV telemetry는 구현되어 있으나 보드별 A/°C 보정과 software protection threshold는 실측 전 사용하지 않는다."),
]
for prefix, text in hardware_items:
    add_bullet(doc, f"{prefix}: {text}", bullet_num_id, bold_prefix=f"{prefix}: ")

add_subheading(doc, "통신·Fault reset 안전 전략")
protocol_rows = [
    ("정상 명령", "20 Hz full frame + 250 ms TTL", "B가 전체 채널·범위·sequence 확인 후 적용"),
    ("TabUI 단절", "A의 AUTO 정책은 계속 수행", "A→B heartbeat는 UI lease와 분리"),
    ("A→B 단절", "B timeout safe-off", "이전 command lease 재사용 금지"),
    ("E-Stop", "falling edge 전체 즉시 차단", "10회 연속 LOW에서 reset-required latch"),
    ("Power Stage Fault", "falling edge 해당 채널 즉시 차단", "20회 LOW 또는 5초 내 3번째 event에서 latch"),
    ("Fault reset", "boot/session/challenge 일치", "B control_result와 1,500 ms correlation 필수"),
]
add_table(doc, ["상황", "동작", "검증 조건"], protocol_rows,
          [1700, 3200, table_width - 4900], font_size=9.2, first_col_bold=True)

add_subheading(doc, "유사 방식과의 비교 및 차별성")
comparison_rows = [
    ("고정 선팅", "상황별 투과율 변경 불가", "센서·모드에 따라 4채널 목표 MI를 연속 변경"),
    ("수동 가림막", "사용자 조작과 일괄 차단 중심", "차박·주차·열부하 시나리오와 AUTO 복귀"),
    ("상용 ON/OFF PDLC", "단일·이진 동작에 제한", "자체 SPWM 전력부와 채널별 중간 산란 제어"),
    ("카메라 HDR/AE", "센서 내부 처리 이후의 보정", "PDLC를 센서 전단 광학 보완 계층으로 활용"),
    ("단일 MCU 구성", "센서 정책과 출력 안전이 결합될 수 있음", "ESP32_A 정책과 ESP32_B 출력·차단 책임 분리"),
]
add_table(doc, ["구분", "기존 방식의 한계", "KUGLASS의 차별성"], comparison_rows,
          [1600, 3600, table_width - 5200], font_size=9.2, first_col_bold=True)

add_subheading(doc, "현재 개발 현황")
status_colors = {
    "완료": "E2F0D9",
    "완료(비실기)": "E2F0D9",
    "제작 완료": "E2F0D9",
    "부분 완료": "FFF2CC",
    "실기 필요": "FCE4D6",
    "미실측": "F4CCCC",
}
status_rows = [
    ("시스템·프로토콜", "완료", "TabUI↔A↔B 책임, v=1 JSON Lines, MI 0.0~0.60 계약 정리", "문서·host test"),
    ("TabUI", "완료(비실기)", "LIVE/MOCK/REPLAY, HMI, 관리자 콘솔, USB gateway, 카메라 viewer", "typecheck·58 tests·build PASS"),
    ("ESP32_A", "완료(비실기)", "카메라·DS18B20·정책·MI servo·20 Hz heartbeat·reset correlation", "제품 host test PASS"),
    ("ESP32_B", "완료(비실기)", "strict parser·4채널 SPWM·E-Stop/Fault/TTL·ADC raw/mV", "제품 host test PASS"),
    ("Logic Carrier", "제작 완료", "1장 as-built 기준, J7·enable gate·Fault pull-up·ADC filter", "hardware contract PASS"),
    ("Power Stage", "제작 완료", "동일 단일 채널 보드 4장", "회로·PCB 원본 확인"),
    ("A–B 실기 통신", "실기 필요", "외부 3선 harness, bridge contention, 장시간 오류율", "10분 HIL 기록 없음"),
    ("ADC·Fault 계측", "부분 완료", "raw/mV 수집과 latch 로직 구현", "보드별 A/°C 보정·trip 실측 없음"),
    ("PDLC·광학·열 통합", "미실측", "4채널 출력, LUT, 영상·열 개선 KPI", "저전압→HV→PDLC 단계 필요"),
]
add_table(doc, ["항목", "상태", "현재 결과", "근거·남은 조건"], status_rows,
          [1700, 1500, 3700, table_width - 6900], font_size=8.7,
          status_col=1, status_colors=status_colors, first_col_bold=True)

add_subheading(doc, "예상 장애요인 및 해결 방안")
risk_rows = [
    ("광학", "PDLC 산란이 포화와 함께 Edge·대비를 낮출 수 있음", "포화율과 Edge Density를 함께 기록하고 MI–Vrms–투과도 LUT와 과차광 억제 기준을 실측"),
    ("전력", "용량성 부하의 과도전류·ringing·발열", "한 보드씩 current-limited 저전압 dummy load 시험 후 네 채널·72 V·PDLC 순으로 확대"),
    ("통신", "UART ROM banner·stale frame·bridge contention", "ROM banner 분류, sequence/TTL 검증, 외부 3선 연속성·10분 오류율 계측"),
    ("Fault", "자기 복구 pulse가 반복 재인가를 만들 수 있음", "첫 두 pulse는 HIGH 안정화 후 복구, 5초 내 세 번째 event 또는 20회 LOW에서 latch"),
    ("ADC", "divider/clamp 부재와 보드별 오차", "허용 범위 확인, 0 V·open·short 경계 시험, 보드별 slope/offset·NTC 계수 산출"),
    ("자료", "Logic Carrier ADC 하위시트·BOM·보드 serial 자료 누락", "실물을 기준으로 식별표를 만들고 회로 원본·계약·시험 기록을 함께 관리"),
    ("네트워크", "현재 plain HTTP·무인증 구성", "localhost 또는 신뢰된 격리 LAN에서만 사용하고 외부 공개 전 TLS·인증·Origin/CSRF 설계"),
]
add_table(doc, ["구분", "장애요인", "해결 방안"], risk_rows,
          [1250, 3450, table_width - 4700], font_size=8.8, first_col_bold=True)

add_subheading(doc, "활용 분야와 발전 방향")
for text in [
    "차량 모형 기반 능동 디밍 글라스 제어 교육·연구 플랫폼, 카메라 전단 광학 보완 실험, 열부하·프라이버시 시나리오 검증, 다채널 PDLC 전력제어 레퍼런스로 활용할 수 있다.",
    "향후에는 실측 데이터가 확보된 뒤 4채널보다 세분화한 window matrix, 경량 예측 모델, 곡면·면적 스케일업을 검토할 수 있다. 현 보고서에서는 이를 현재 구현으로 표현하지 않는다.",
]:
    add_bullet(doc, text, bullet_num_id)

# Detailed operation and equipment.
add_section_heading(doc, "작품 상세 설명 및 지원 장비 사용 계획")
add_subheading(doc, "작품 작동 원리")
flow_rows = [
    ("1. 입력 수집", "OV2640 영상과 DS18B20 내부온도를 ESP32_A가 취득한다. 카메라 frame timestamp, 온도 CRC·범위·stale을 함께 관리한다."),
    ("2. 영상·온도 처리", "정방향 frame의 운전석측/조수석측 ROI에서 밝기·포화·highlight·Edge Density를 계산하고 thermal risk 입력을 구성한다."),
    ("3. 목표 산출", "상황 모드, privacy, thermal, camera glare를 결합해 CH0~CH3 target_mi와 enable을 계산한다."),
    ("4. 명령 안정화", "ESP32_A가 AUTO deadband와 policy servo를 적용해 commanded_mi를 만든다. target과 commanded를 구분해 기록한다."),
    ("5. A→B 전달", "UART1 115200 8-N-1 JSON Lines로 CH0~CH3 full frame을 20 Hz 전송한다. seq와 ttl_ms=250을 포함한다."),
    ("6. 실시간 출력", "ESP32_B가 frame과 안전 입력을 검증한 뒤 최종 slew, 16 kHz SPWM, 60 Hz polarity와 1 ms 방향 blanking을 적용한다."),
    ("7. 물리 구동", "Logic Carrier의 enable gate와 J7을 거쳐 각 Power Stage가 H-Bridge·LC filter로 PDLC 출력 파형을 만든다."),
    ("8. 상태 피드백", "ESP32_B가 applied_mi·Fault·E-Stop·ADC·reset 결과를 A로 보내고 A가 TabUI에 target/commanded/applied 상태를 중계한다."),
]
add_table(doc, ["단계", "내용"], flow_rows,
          [1700, table_width - 1700], font_size=9.2, first_col_bold=True)

add_subheading(doc, "주요 시연 시나리오")
demo_rows = [
    ("기본/주행", "카메라·내부온도 기반 AUTO", "실제 센서 quality, 판단 이유, 네 채널 상태 표시"),
    ("카메라 역광", "강한 광원 전·후 ROI 지표 비교", "포화 완화와 Edge 보존을 동시에 평가"),
    ("열부하", "DS18B20 온도 상승에 따른 MI 변화", "MOCK/HIL 합성값은 물리 센서와 명확히 구분"),
    ("차박·주차", "전 채널 enable=true의 비영 MI 강산란", "safe-off MI 0과 프라이버시 동작을 구분"),
    ("수동·AUTO", "일반 15초 TTL·관리자 지속·return_auto", "제어 출처와 남은 TTL 표시"),
    ("Fault·단절", "E-Stop, 채널 Fault, A/B stale", "전체/채널 safe-off와 명령 차단 확인"),
]
add_table(doc, ["시나리오", "입력·동작", "관찰 항목"], demo_rows,
          [1700, 3400, table_width - 5100], font_size=9.2, first_col_bold=True)

add_subheading(doc, "현재 비실기 검증 결과")
validation_rows = [
    ("하드웨어 계약", "PASS", "원본 hash, GPIO, J7/J10, ESP32_B pin/ADC descriptor 정합"),
    ("TabUI check", "PASS", "TypeScript typecheck + Python unit test 58개"),
    ("TabUI build", "PASS", "Vite production build, 1,615 modules transformed"),
    ("ESP32_A host test", "PASS", "카메라·온도·정책·프로토콜·reset·telemetry·프로젝트 독립성"),
    ("ESP32_B host test", "PASS", "strict parser·4채널 출력·slew·E-Stop/Fault·ADC·status"),
    ("A–B BAD_JSON", "PASS", "ROM boot noise와 ADC 포함 status의 cross-project 호환"),
]
add_table(doc, ["검증", "결과", "확인 범위"], validation_rows,
          [2100, 1100, table_width - 3200], font_size=9.2, first_col_bold=True,
          status_col=1, status_colors={"PASS": "E2F0D9"})
add_callout(doc, "검증 해석", "위 결과는 소프트웨어·문서·정적 계약의 통과 결과다. Power Stage/PDLC/HV의 실제 파형·발열·절연·광학 성능 통과를 의미하지 않는다.", fill="FFF7E6", border="D8B56B")

add_subheading(doc, "지원 장비 사용 방안")
equipment_rows = [
    ("태블릿·MacBook", "HMI 표시, 명령 입력, LIVE 상태 기록", "신뢰된 격리 LAN 또는 localhost에서 사용"),
    ("오실로스코프", "16 kHz carrier, 60 Hz polarity, blanking, Fault latency, bridge/output 파형", "절연·정격에 맞는 probe와 접지 방식 사용"),
    ("멀티미터·연속성계", "J7/J10 pin 1, odd/even GND, harness, 전원 극성 확인", "무전원 단계에서 우선 수행"),
    ("가변 전원·dummy load", "logic-only·저전압 Power Stage 단계 시험", "current limit 설정 후 한 채널씩 확대"),
    ("기준 온도계·전류계", "DS18B20과 보드별 ADC slope/offset·NTC 보정", "조건·장비·보드 식별을 시험 기록에 포함"),
    ("고휘도 광원·열원", "역광·열부하 시나리오 재현", "거리·각도·시간을 고정해 전후 비교"),
    ("안전 외함·표시", "72 V·PDLC 단계 접근 통제와 방전 확인", "E-Stop이 rail 차단이 아님을 전제로 운영"),
]
add_table(doc, ["장비", "사용 목적", "운영 원칙"], equipment_rows,
          [2000, 3900, table_width - 5900], font_size=8.9, first_col_bold=True)

add_subheading(doc, "실기 검증 순서")
hil_steps = [
    "무전원 상태에서 U3 방향, J7/J10 pin 1, 모든 짝수 GND, 전원 극성과 A–B 3선 harness 연속성을 확인한다.",
    "Power Stage와 72 V를 분리한 Logic Carrier 저전압 조건에서 부팅·reset의 ENABLE LOW/PWM force-low, E-Stop과 FAULT_N pulse를 확인한다.",
    "Power Stage를 한 장씩 연결해 RUN_OK, 16 kHz/60 Hz, direction blanking과 정상 MI slew를 logic-only·저전압 dummy load에서 검증한다.",
    "8개 ADC의 0 V·정상·open·short 경계를 확인하고 네 보드별 current slope/offset과 NTC 변환 계수를 기록한다.",
    "앞 단계 PASS 보드만 current-limited DC bus와 dummy load, 네 채널, PDLC 순으로 확대하고 광학·열 KPI와 안전 증거를 남긴다.",
]
for step in hil_steps:
    add_bullet(doc, step, bullet_num_id)

# Current progress and schedule.
add_section_heading(doc, "개발 현황 및 잔여 일정", page_break=True)
add_subheading(doc, "현재 기준선")
add_body(doc, "2026년 8월 10일 기준으로 KUGLASS의 제품 구조, TabUI, ESP32_A/ESP32_B canonical firmware, 통신 계약, 제작 하드웨어 계약과 비실기 회귀 검증이 정리되었다. 다음 단계의 핵심은 제작된 네 Power Stage를 실기에서 개별 식별·계측하고, 저전압부터 통합하여 광학·열 성능을 수치로 닫는 것이다.")

schedule_rows = [
    ("0", "현재 기준선", "완료", "비실기 계약·host test·TabUI build", "재현 가능한 PASS 기록"),
    ("1", "예선 이후 1주", "예정", "무전원 연속성·Logic Carrier 저전압·UART 10분 HIL", "pin 1·전원·E-Stop·Fault·통신 PASS"),
    ("2", "예선 이후 2주", "예정", "Power Stage 1장씩 logic-only·저전압 시험", "CH0~CH3 carrier·blanking·RUN_OK PASS"),
    ("3", "예선 이후 3~4주", "예정", "ADC 입력 범위·보드별 보정·4채널 동시 시험", "raw/mV validity와 calibration record"),
    ("4", "예선 이후 5주", "예정", "PDLC 연결, MI sweep, 카메라·열부하 KPI 측정", "MI–Vrms–투과도 LUT와 전후 비교"),
    ("5", "예선 이후 6주", "예정", "1:10 모형·TabUI 통합, 4개 시나리오 리허설", "시연 체크리스트·Fault 복구 기록"),
    ("6", "결선 제출 전", "예정", "영상·보고서·시험 evidence 정리", "재현 가능한 시연·최종 제출물"),
]
add_table(doc, ["단계", "시점", "상태", "주요 작업", "완료 기준"], schedule_rows,
          [650, 1750, 1050, 3650, table_width - 7100], font_size=8.7,
          status_col=2, status_colors={"완료": "E2F0D9", "예정": "FFF2CC"}, first_col_bold=True)

add_subheading(doc, "단계별 판단 원칙")
for text in [
    "이전 단계의 정적 검사나 안전 입력 검증이 실패하면 출력 전압을 높이거나 다음 보드를 추가하지 않는다.",
    "네 Power Stage 각각의 식별 정보, 시험 조건, 계측기, 측정값과 사진·scope capture를 기록한다.",
    "회로 명목값 또는 host test 결과를 실측 성능으로 대체하지 않으며, 미확정 값은 보고서와 UI에서 계속 미실측으로 표시한다.",
]:
    add_bullet(doc, text, bullet_num_id)

# Team and sources.
add_section_heading(doc, "팀 구성 및 역량", page_break=True)
team_rows = [
    ("팀장", "최무영", ["프로젝트 일정·문서·발표 총괄", "시스템 아키텍처와 TabUI/ESP32_A 통합", "카메라 지표·HMI·상태 모델 검증"], "Python, TypeScript/React, C, OpenCV, 웹·AI 프로젝트와 시스템 통합 경험"),
    ("팀원", "전수찬", ["1:10 차량 모형·제어 외함 설계·조립", "웹 HMI용 3D 모델과 시연 환경 제작", "배선·하네스·시연 영상 구성"], "Fusion 360, Blender, Python, 3D 프린팅·기구 조립·배선 경험"),
    ("팀원", "김덕렬", ["ESP32_A/B 실시간 펌웨어와 host test", "A–B protocol, SPWM, Fault·TTL 동작 검증", "logic-only·저전압 HIL 수행"], "C/C++, ESP-IDF, FreeRTOS, MCPWM/PWM, PSIM·PSpice 전력제어 경험"),
    ("팀원", "이승민", ["Logic Carrier·Power Stage 회로/PCB 기준 관리", "전원·안전 회로와 계측 절차 수행", "출력 파형·발열·ADC·Fault KPI 분석"], "KiCad, Altium, LTspice, PSIM, Python 데이터 분석과 전력전자 설계 경험"),
]
team_table = doc.add_table(rows=1, cols=4)
configure_table_geometry(team_table, [1050, 1250, 3900, table_width - 6200])
set_repeat_table_header(team_table.rows[0])
prevent_row_split(team_table.rows[0])
for i, header in enumerate(["구분", "성명", "현재 담당 업무", "업무 관련 역량"]):
    cell = team_table.rows[0].cells[i]
    set_cell_shading(cell, "E5E5E5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(header)
    set_run_font(r, FONT_BODY, 9.3, True)
for role, name, duties, skills in team_rows:
    row = team_table.add_row()
    prevent_row_split(row)
    for index, value in enumerate((role, name)):
        cell = row.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, FONT_BODY, 9.0, True if index == 0 else False)
    duty_cell = row.cells[2]
    duty_cell.text = ""
    for idx, duty in enumerate(duties):
        p = duty_cell.paragraphs[0] if idx == 0 else duty_cell.add_paragraph()
        apply_bullet(p, bullet_num_id, 0)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(duty)
        set_run_font(r, FONT_BODY, 8.8)
    skill_cell = row.cells[3]
    skill_cell.text = ""
    p = skill_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(skills)
    set_run_font(r, FONT_BODY, 8.8)
    for index, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, [1050, 1250, 3900, table_width - 6200][index])
        set_cell_margins(cell, top=100, bottom=100)
doc.add_paragraph(style="KUG Caption")

add_section_heading(doc, "참고 문헌 및 기준 자료")
references = [
    "KUGLASS_DEV, README.md, 개발 계획서.md, docs/ARCHITECTURE.md, docs/PROTOCOL.md, docs/VALIDATION.md (2026.08.10 기준).",
    "KUGLASS_DEV, TabUI/README.md 및 TabUI tests; ESP32_A_Algo/README.md·host_tests; ESP32_B_Algo/README.md·host_tests.",
    "KUGLASS_DEV, hardware/manifest.json, hardware/contracts/esp32_b_io.json, power_stage.json, safety.json.",
    "KUGLASS Logic Carrier, Logic carrier.pdf, 2 pages, as-built schematic; KUGLASS Power Stage, Power_stage.pdf 및 KiCad schematic/PCB.",
    "Espressif Systems, ESP-IDF Programming Guide (프로젝트 기준 6.0.2) 및 ESP32-S3-DevKitC-1 User Guide.",
    "Espressif Systems, esp32-camera component 2.1.7.",
    "DS18B20 Programmable Resolution 1-Wire Digital Thermometer Datasheet.",
    "IRS2104(S) High and Low Side Driver Datasheet; TLV1701 Comparator Datasheet; 74LVC04/74LVC08/74HC08 Datasheets (저장소 hardware/datasheets 수록본).",
]
for index, item in enumerate(references, 1):
    p = doc.add_paragraph(style="KUG Reference")
    p.paragraph_format.left_indent = Twips(360)
    p.paragraph_format.first_line_indent = Twips(-360)
    r1 = p.add_run(f"[{index}] ")
    set_run_font(r1, FONT_BODY, 9.5, True)
    r2 = p.add_run(item)
    set_run_font(r2, FONT_BODY, 9.5)

add_callout(doc, "제출 기준", "본 보고서는 저장소의 현재 구현과 2026.08.10 비실기 검증 결과를 기준으로 작성했으며, 실기 계측이 없는 값은 확정 성능으로 제시하지 않았다.", fill="EEF7F0", border="A8C9B0")

# Prevent a final empty body paragraph from creating a stray page.
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.widow_control = True

doc.save(OUTPUT)
print(OUTPUT)
